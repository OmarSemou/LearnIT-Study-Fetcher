from __future__ import annotations

from pathlib import Path

from docx import Document


def extract_text(path: str | Path) -> str:
    document = Document(str(path))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    table_text: list[str] = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_text.append(" | ".join(cells))
    return "\n".join([*paragraphs, *table_text])
